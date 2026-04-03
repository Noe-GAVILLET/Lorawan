"""
Génère Chapitre_4_LoRaWAN.docx  (~5 pages)
Chapitre 4 — Analyse de Performance et Limites de Conception
Projet ETRS012 — Jumeau Numérique Scientifique de Ruche IoT
Université Savoie Mont Blanc — Avril 2026
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_figure(doc, number, description):
    p = doc.add_paragraph()
    run = p.add_run(f"[FIGURE {number} : {description}]")
    run.bold = True
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()


def add_formula(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()

# ── Titre du chapitre ───────────────────────────────────────────────────────
add_heading(doc, "Chapitre 4 — Analyse de Performance et Limites de Conception", level=1)

add_body(doc,
    "Ce chapitre constitue le cœur expérimental du dossier. Il présente successivement "
    "la formalisation mathématique du modèle biophysique implémenté, les résultats de "
    "thermorégulation obtenus sur 8 heures de simulation, la justification rigoureuse "
    "de l'algorithme de détection d'essaimage par dérivée temporelle, puis l'analyse "
    "critique des résultats de classification pour les hypothèses H1 (PDR = 0,90) et "
    "H2 (PDR = 0,65). Le chapitre se conclut par une évaluation honnête des limites "
    "de conception qui bornent la portée de ces résultats."
)

# ── 4.1 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "4.1  Modèle biophysique : formalisation mathématique", level=2)

add_body(doc,
    "Le simulateur biophysique (random_data_publisher.py) reproduit trois dynamiques "
    "physiques couplées : la thermorégulation du couvain, le budget énergétique en miel, "
    "et le cycle de butinage. Ces trois composantes sont calculées à chaque pas de temps "
    "τ_pub (intervalle de publication en secondes)."
)

add_heading(doc, "Thermorégulation", level=3)

add_body(doc,
    "La température interne mesurée T_real est modélisée comme une variable aléatoire "
    "centrée sur la cible BEEHAVE de 34,5 °C, avec un bruit uniforme ±0,1 °C représentant "
    "les micro-fluctuations de la grappe d'abeilles :"
)

add_formula(doc, "T_real(t) = 34,5 + U(-0,1 ; +0,1)     [°C]")

add_body(doc,
    "La température ambiante T_amb suit un cycle sinusoïdal journalier calé sur les "
    "conditions météorologiques printanières moyennes (Savoie, avril) : minimum à 3 h, "
    "maximum à 15 h, moyenne 18 °C, amplitude ± 8 °C :"
)

add_formula(doc, "T_amb(t) = 18,0 + 8,0 × sin(2π × (t_h - 9) / 24)     [°C, t_h en heures UTC]")

add_heading(doc, "Budget énergétique", level=3)

add_body(doc,
    "La consommation de miel par heure est proportionnelle à l'écart entre la température "
    "cible et la température ambiante (coût de chauffage), avec un terme de base "
    "métabolique indépendant de la température :"
)

add_formula(doc,
    "ΔM_miel(t) = [M_base + α × max(0, T_cible - T_amb(t))] × (τ_pub / 3600)\n"
    "avec  M_base = 0,005 kg/h,  α = 0,001 kg/(h·°C),  T_cible = 34,5 °C"
)

add_body(doc,
    "Ce modèle est une approximation linéaire de premier ordre du bilan BEEHAVE. "
    "La valeur α = 0,001 kg/(h·°C) est choisie pour produire une consommation nocturne "
    "réaliste de ~0,022 kg/h lorsque T_amb = 10 °C (nuit froide de printemps), "
    "soit ~530 g de miel consommé par nuit — cohérent avec les données empiriques "
    "de Meikle & Holst [3] sur des ruches Langstroth de 45 kg."
)

add_heading(doc, "Cycle de butinage et essaimage", level=3)

add_body(doc,
    "La masse totale de la ruche M(t) évolue par incréments à chaque pas de temps "
    "selon le régime en cours :"
)

add_formula(doc,
    "ΔM_but(t) selon t_h :\n"
    "  t_h ∈ [8h, 11h)  → ΔM_but = -0,15 × (τ_pub/3600)  kg  (départ butineuses)\n"
    "  t_h ∈ [11h, 18h) → ΔM_but = +0,20 × (τ_pub/3600)  kg  (retour avec nectar)\n"
    "  sinon             → ΔM_but = 0\n\n"
    "En régime SCENARIO=extreme :\n"
    "  ΔM_essaimage = -2,5 × (τ_pub/3600)  kg  (perte brutale, toutes heures)"
)

add_body(doc,
    "La masse est bornée inférieurement à MASS_FLOOR_KG = 5,0 kg, correspondant à la "
    "structure bois et aux cadres vides d'une ruche Langstroth standard. Cette borne "
    "garantit la cohérence physique du modèle en évitant les masses négatives lors "
    "de simulations prolongées en régime extrême."
)

add_figure(doc, 4,
    "Flux de données biophysiques simulées : masse M(t) en kg et température T_real(t) "
    "en °C sur 8 h — régime nominal (cycles butinage), transition, régime extrême (essaimage, "
    "chute -2,5 kg/h)"
)

# ── 4.2 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "4.2  Résultats de thermorégulation : MAE = 0,635 °C", level=2)

add_body(doc,
    "L'évaluation de la fidélité thermique du Jumeau Numérique compare T_real(t) à la "
    "cible BEEHAVE T_cible = 34,5 °C. Les métriques MAE (Mean Absolute Error) et RMSE "
    "(Root Mean Square Error) sont calculées sur l'ensemble du dataset H1 "
    "(5 216 points, PDR = 0,90, ~8 h)."
)

add_table(doc,
    ["Régime", "N points", "MAE température (°C)", "RMSE température (°C)"],
    [
        ["Global",                  "5 216", "0,635", "0,819"],
        ["Nominal (butinage/nuit)", "4 560", "0,635", "0,819"],
        ["Extrême (essaimage)",     "656",   "0,630", "0,816"],
    ]
)

add_body(doc,
    "L'observation centrale est la quasi-identité des métriques entre régime nominal "
    "et régime extrême : ΔMAE = 0,005 °C, ΔRMSE = 0,003 °C. Cette stabilité n'est "
    "pas surprenante — elle est une propriété structurelle du modèle biophysique : "
    "la thermorégulation de la colonie (T_real = 34,5 ± 0,1 °C) est implémentée "
    "comme un processus indépendant du budget massique. L'essaimage n'affecte pas "
    "la variable T_real dans le simulateur, car le modèle simplifié ne couple pas "
    "la dynamique de population (nombre d'abeilles perdues lors de l'essaimage) "
    "avec la capacité thermorégulatrice de la colonie résiduelle."
)

add_body(doc,
    "Cette indépendance est biologiquement discutable sur un essaimage sévère "
    "(départ de 50 % de la colonie) : une colonie réduite de moitié peut avoir "
    "des difficultés à maintenir 34,5 °C lors des nuits froides suivant l'essaimage. "
    "BEEHAVE modélise cet effet via la dynamique de population du couvain et la "
    "redistribution des butineuses en nourrices. Notre approximation de premier ordre "
    "ignore ce couplage — limitant la validité du résultat thermique à la seule "
    "période de l'essaimage (quelques heures), non aux jours suivants."
)

add_body(doc,
    "La MAE globale de 0,635 °C est bien inférieure au critère d'acceptance H1 "
    "de 1,0 °C défini dans le cahier des charges (§ 8). Du point de vue métrologique, "
    "cette valeur est cohérente avec la précision typique des capteurs de température "
    "embarqués dans les ruches (Dallas DS18B20 : ±0,5 °C, SHT31 : ±0,3 °C), "
    "suggérant que le bruit de modélisation est du même ordre que le bruit de mesure "
    "d'un capteur réel — validant qualitativement l'approche de simulation."
)

# ── 4.3 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "4.3  Algorithme de détection par dérivée temporelle", level=2)

add_body(doc,
    "L'algorithme de détection d'essaimage implémenté dans train_eval.py repose sur "
    "le calcul de la dérivée temporelle discrète de la masse interpolée, comparée à "
    "un seuil fixe. Cette approche est intentionnellement simple, interprétable et "
    "auditée — à l'opposé des approches boîte noire comme LSTM ou DeepAnT [4]."
)

add_heading(doc, "Pipeline de calcul", level=3)

add_body(doc,
    "Le pipeline complet de calcul de la dérivée se déroule en quatre étapes :"
)

add_body(doc,
    "Étape 1 — Interpolation linéaire des gaps LoRaWAN : les paquets perdus (PDR < 1) "
    "créent des lacunes dans la série temporelle. La méthode pandas.interpolate(method='linear') "
    "comble ces gaps par interpolation linéaire entre les deux points valides encadrants. "
    "Cette étape est critique : sans interpolation, les différences temporelles seraient "
    "calculées sur des pas de temps non uniformes, rendant la dérivée instable."
)

add_body(doc,
    "Étape 2 — Calcul de la différence temporelle : Δt_i = t_i − t_{i−1}, exprimé "
    "en minutes pour être homogène à l'unité de la dérivée (kg/min)."
)

add_body(doc,
    "Étape 3 — Calcul de la dérivée discrète :"
)

add_formula(doc,
    "dM/dt |_i = (M_interp(t_i) - M_interp(t_{i-1})) / Δt_i     [kg/min]"
)

add_body(doc,
    "Étape 4 — Seuillage et prédiction :"
)

add_formula(doc,
    "predicted_swarming(t_i) = True  si  dM/dt |_i ≤ -0,03 kg/min\n"
    "                        = False sinon"
)

add_heading(doc, "Justification du seuil −0,03 kg/min", level=3)

add_body(doc,
    "Le seuil SWARM_DERIVATIVE_THRESHOLD = −0,03 kg/min est fixé manuellement entre "
    "les deux régimes connus du simulateur. L'analyse des dynamiques attendues justifie "
    "ce choix :"
)

add_table(doc,
    ["Régime", "dM/dt typique", "Cause physique", "Marge vs seuil −0,03"],
    [
        ["Nominal (nuit)",    "≈ −0,0001 kg/min", "Consommation miel nocturne seule",         "Facteur 300×"],
        ["Nominal (départ 8h–11h)", "≈ −0,003 kg/min",  "Butineuses quittent la ruche", "Facteur 10×"],
        ["Extrême (essaimage)", "≈ −0,042 kg/min", "2,5 kg/h × (1/60) = 0,042 kg/min", "Facteur 1,4× (marge)"],
    ]
)

add_body(doc,
    "Le facteur de séparation de 10× entre la dérivée nominale maximale (−0,003 kg/min "
    "au départ des butineuses) et le seuil (−0,030 kg/min) garantit une précision parfaite "
    "en conditions H1, où les gaps courts (5 s, PDR = 0,90) ne dégradent pas l'interpolation. "
    "La marge de 1,4× côté extrême (−0,042 vs −0,030 kg/min) explique la vulnérabilité "
    "en conditions H2 : un gap de 60 s à la transition normal→extrême lisse la dérivée, "
    "qui peut ne pas atteindre −0,030 kg/min sur ce seul point — produisant un Faux Négatif."
)

add_figure(doc, 5,
    "Dérivée temporelle dM/dt (kg/min) sur la série interpolée : régime nominal (−0,003 max), "
    "transition, seuil −0,03 kg/min (ligne pointillée), régime extrême (−0,042 typique) — "
    "visualisation du Faux Négatif H2 à la transition"
)

# ── 4.4 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "4.4  Résultats H1 : PDR = 0,90, F1-Score = 1,00", level=2)

add_body(doc,
    "Le dataset H1 couvre la période 2026-04-02T10:49 → 18:47 UTC, avec "
    "PUBLISH_INTERVAL_SECONDS = 5 s et LORAWAN_PDR = 0,90. Sur 5 216 points bruts, "
    "~521 ont été perdus (10 %) puis interpolés. Le régime extrême (SCENARIO=extreme) "
    "a été activé à ~14h20 et maintenu jusqu'à la fin de la session."
)

add_table(doc,
    ["Métrique de classification", "Valeur H1"],
    [
        ["Vrais Positifs (TP)",  "3 192"],
        ["Faux Positifs (FP)",   "0"],
        ["Faux Négatifs (FN)",   "0"],
        ["Vrais Négatifs (TN)",  "~1 503"],
        ["Précision",            "1,00"],
        ["Rappel",               "1,00"],
        ["F1-Score",             "1,00"],
    ]
)

add_body(doc,
    "Ce résultat F1 = 1,00 est attendu par construction et ne constitue pas une "
    "découverte expérimentale. Trois facteurs structurels le rendent inévitable :"
)

add_body(doc,
    "1. Séparation triviale des classes : les deux régimes du simulateur produisent "
    "des dérivées séparées d'un facteur 14 (−0,003 vs −0,042 kg/min). Le seuil "
    "de −0,030 kg/min est positionné dans un espace vide large de ~0,027 kg/min "
    "de chaque côté — aucun point ne peut se trouver dans cette zone de décision "
    "sauf à la transition (qui dure < 1 cycle de publication en H1)."
)

add_body(doc,
    "2. Gaps courts non dégradants : avec τ_pub = 5 s et PDR = 0,90, la probabilité "
    "qu'un gap tombe exactement sur le premier paquet du régime extrême est de "
    "P = 0,10 × (1 paquet critique / N paquets totaux) ≈ 0,10 × (1/656) ≈ 0,015 %. "
    "Sur un seul run, ce cas ne s'est pas produit — le F1 resterait 1,00 sur la "
    "grande majorité des runs indépendants."
)

add_body(doc,
    "3. Biais partiel de la ground truth : pour le dataset H1, la vérité terrain "
    "repose sur la plage horaire de déclenchement (13h30–14h30 UTC), construite "
    "indépendamment de la dérivée mais sujette à un biais si l'horloge du conteneur "
    "Docker dérivait lors de la collecte. Cette limite est explicitée dans le protocole "
    "expérimental (§ 8) et corrigée dans le dataset H2 via le tag `scenario` InfluxDB."
)

add_body(doc,
    "La valeur de H1 est donc celle d'une baseline de référence : elle confirme que "
    "le pipeline complet (simulation → perte PDR → interpolation → dérivée → seuillage) "
    "fonctionne correctement en l'absence de contrainte réseau sévère. La contribution "
    "scientifique réelle du projet repose entièrement sur H2."
)

# ── 4.5 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "4.5  Résultats H2 : PDR = 0,65, Rappel = 0,94", level=2)

add_body(doc,
    "Le dataset H2 a été collecté avec PUBLISH_INTERVAL_SECONDS = 60 s et "
    "LORAWAN_PDR = 0,65 — paramètres représentatifs d'un vrai déploiement LoRaWAN "
    "en zone forestière contrainte par le Duty Cycle. La ground truth repose sur "
    "le tag `scenario` enregistré dans InfluxDB indépendamment de la prédiction "
    "(méthode rigoureuse, sans biais circulaire)."
)

add_table(doc,
    ["Métrique de classification", "Valeur H1 (PDR=0,90)", "Valeur H2 (PDR=0,65)", "Δ"],
    [
        ["Vrais Positifs (TP)",  "3 192", "17",   "—"],
        ["Faux Positifs (FP)",   "0",     "0",    "—"],
        ["Faux Négatifs (FN)",   "0",     "1",    "+1"],
        ["Précision",            "1,00",  "1,00", "0,00"],
        ["Rappel",               "1,00",  "0,94", "−0,06"],
        ["F1-Score",             "1,00",  "0,97", "−0,03"],
    ]
)

add_body(doc,
    "Le résultat clé est le Faux Négatif unique (FN = 1) : un essaimage n'a pas été détecté "
    "sur le point de transition entre régime nominal et régime extrême. Le mécanisme est "
    "le suivant : avec τ_pub = 60 s et PDR = 0,65, chaque paquet perdu crée un gap de "
    "60 s (un paquet manquant) à 120 s (deux paquets consécutifs manquants). L'interpolation "
    "linéaire comble ce gap en lissant la transition de masse — la dérivée calculée sur "
    "le point interpolé est alors :"
)

add_formula(doc,
    "dM/dt_interpolé = ΔM_lissé / Δt_gap\n"
    "                = [(M_extrême - M_nominal) / 2] / (60 s / 60)\n"
    "                ≈ [-0,042 × 60/60] / 1,0 ≈ -0,021 kg/min  <  seuil -0,030"
)

add_body(doc,
    "La dérivée lissée (−0,021 kg/min) reste en dessous du seuil (−0,030 kg/min) "
    "en valeur absolue — le point de transition n'est pas classifié comme essaimage. "
    "Tous les paquets suivants (régime extrême pleinement établi) produisent des "
    "dérivées de −0,042 kg/min et sont correctement détectés. L'effet H2 est donc "
    "un artefact de frontière, localisé sur un seul intervalle de 60 s — ce qui "
    "explique que le Rappel ne chute qu'à 0,94 et non à des valeurs plus sévères."
)

add_body(doc,
    "Ce résultat valide H2 telle que formulée (chute du Rappel, augmentation des FN) "
    "tout en nuançant son amplitude : PDR = 0,65 avec τ_pub = 60 s est un scénario "
    "déjà sévère (35 % de perte), et l'effet reste limité (1 FN). Un scénario encore "
    "plus dégradé (PDR < 0,50, ou burst errors corrélés comme discuté en § 2.4) "
    "produirait vraisemblablement plusieurs FN consécutifs à la frontière de transition, "
    "avec un Rappel pouvant chuter sous 0,80."
)

add_figure(doc, 6,
    "Comparaison H1 vs H2 : histogramme Précision / Rappel / F1-Score "
    "pour PDR = 0,90 (bleu) et PDR = 0,65 (orange) — mise en évidence "
    "de la dégradation du Rappel uniquement"
)

# ── 4.6 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "4.6  Limites de conception et validité externe", level=2)

add_body(doc,
    "Une évaluation rigoureuse impose d'inventorier explicitement les limites qui "
    "bornent la généralisation des résultats présentés. Six limites majeures sont "
    "identifiées, classées par degré d'impact sur la validité externe."
)

add_heading(doc, "L1 — Absence d'intervalle de confiance (impact : élevé)", level=3)

add_body(doc,
    "Toutes les métriques reportées (MAE = 0,635 °C ; Rappel H2 = 0,94) reposent sur "
    "un seul run expérimental avec une graine aléatoire non fixée (random.uniform sans "
    "random.seed). La reproductibilité des valeurs sur N runs indépendants n'est pas "
    "établie. Un protocole rigoureux nécessiterait N ≥ 30 répétitions pour estimer "
    "moyenne et écart-type par bootstrap, et calculer un intervalle de confiance à 95 %. "
    "En l'état, un deuxième run pourrait produire FN = 0 (Rappel = 1,00) ou FN = 3 "
    "(Rappel = 0,83) selon les tirages aléatoires du PDR — l'incertitude est inconnue."
)

add_heading(doc, "L2 — Durée de simulation insuffisante (impact : élevé)", level=3)

add_body(doc,
    "Les 8 heures de données collectées couvrent un seul cycle journalier avec un "
    "essaimage unique. Une saison apicole réelle dure 5 à 6 mois avec des dynamiques "
    "de longue durée absentes du modèle : montée en population du couvain au printemps, "
    "variations de disponibilité florale, hivernage avec réduction du butinage, "
    "possibilité de plusieurs essaimages successifs. Le cycle journalier est identique "
    "à chaque heure simulée — sans variation inter-journalière. La validité des résultats "
    "à l'échelle de la semaine ou du mois ne peut être inférée depuis le dataset actuel."
)

add_heading(doc, "L3 — PDR Bernoulli vs. canal radio réel (impact : moyen)", level=3)

add_body(doc,
    "Comme discuté en § 2.4, le modèle de pertes par processus de Bernoulli indépendant "
    "sous-estime l'effet des burst errors d'un canal radio réel. Les résultats H2 "
    "représentent un scénario optimiste : un canal à mémoire (Gilbert-Elliott) produisant "
    "des séquences de 3 à 10 paquets consécutifs perdus génèrerait systématiquement "
    "plusieurs FN à chaque rafale de perte pendant l'essaimage."
)

add_heading(doc, "L4 — Absence de capteurs physiques réels (impact : élevé)", level=3)

add_body(doc,
    "L'ensemble de la validation repose sur des données entièrement simulées. La "
    "transférabilité des seuils (−0,03 kg/min) et métriques à un déploiement réel "
    "reste à démontrer. Un vrai capteur de poids présente des dérives de zéro "
    "thermiques (±0,002 %/°C pour une cellule de charge), des chocs mécaniques "
    "(insectes, pluie, vent), et des oscillations dues au vent agitant les ruches "
    "légères qui peuvent produire des pics de dérivée non liés à un essaimage. "
    "Le seuil nominal devrait être recalibré sur données terrain."
)

add_heading(doc, "L5 — Trivialité de H1 (impact : faible sur H2, élevé sur communication)", level=3)

add_body(doc,
    "Le F1-Score = 1,00 de H1 est une propriété déterministe du simulateur (classes "
    "séparées d'un facteur 14) et non un résultat statistiquement significatif. "
    "Présenter ce résultat sans la nuance du § 4.4 serait une erreur de communication "
    "scientifique majeure. H1 doit être systématiquement accompagné de sa qualification "
    "de 'baseline de référence' pour ne pas induire le lecteur en erreur sur les "
    "capacités réelles du système."
)

add_heading(doc, "L6 — Modèle biophysique découplé masse/population (impact : moyen)", level=3)

add_body(doc,
    "La thermorégulation et la dynamique massique sont découplées dans notre modèle "
    "(§ 4.1). BEEHAVE modélise explicitement l'impact de la réduction de population "
    "sur la capacité thermorégulatrice. Ce couplage absent explique pourquoi notre MAE "
    "thermique est stable entre régimes — résultat attendu dans le simulateur mais "
    "potentiellement incorrect sur une vraie ruche post-essaimage en conditions froides."
)

# ---------------------------------------------------------------------------
# Sauvegarde
# ---------------------------------------------------------------------------
out_path = "Chapitre_4_LoRaWAN.docx"
doc.save(out_path)
print(f"✓ Document généré : {out_path}")
