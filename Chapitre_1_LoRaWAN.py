"""
Génère Chapitre_1_LoRaWAN.docx  (~3 pages)
Chapitre 1 — Introduction et État de l'art
Projet ETRS012 — Jumeau Numérique Scientifique de Ruche IoT
Université Savoie Mont Blanc — Avril 2026
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_figure(doc, number, description):
    p = doc.add_paragraph()
    run = p.add_run(f"[FIGURE {number} : {description}]")
    run.bold = True
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()

# ── Titre du chapitre ───────────────────────────────────────────────────────
add_heading(doc, "Chapitre 1 — Introduction et État de l'Art", level=1)

add_body(doc,
    "Ce chapitre situe le projet dans son contexte scientifique et sociotechnique. "
    "Il présente les motivations qui justifient l'emploi d'un Jumeau Numérique (Digital Twin) "
    "pour la surveillance apicole, formule les questions de recherche et les hypothèses testées, "
    "puis synthétise de manière critique la littérature sur laquelle repose l'architecture retenue."
)

# ── 1.1 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "1.1  Contexte : IoT, LPWAN et apiculture de précision", level=2)

add_body(doc,
    "L'Internet des Objets (IoT) désigne un ensemble hétérogène de dispositifs embarqués "
    "capables de collecter, traiter et transmettre des données physiques vers des systèmes "
    "d'information distants. Son déploiement massif dans les secteurs agricole, industriel et "
    "urbain pose une contrainte fondamentale : la majorité des nœuds capteurs opèrent sur "
    "batterie dans des zones à couverture cellulaire limitée, ce qui exclut les protocoles "
    "à haute consommation comme le Wi-Fi ou la 4G/LTE."
)

add_body(doc,
    "Les réseaux LPWAN (Low-Power Wide-Area Network) ont émergé pour répondre à ce compromis "
    "triples — portée, consommation, coût — inaccessible aux technologies classiques. Parmi eux, "
    "LoRaWAN (Long Range Wide Area Network) s'est imposé comme standard de facto en milieu "
    "agricole et rural grâce à une portée typique de 2 à 15 km, une consommation en veille "
    "inférieure à 10 µA, et une durée de vie de batterie pouvant atteindre 3 à 5 ans selon "
    "l'intervalle de publication. Cette longévité constitue un argument éthique et environnemental "
    "décisif : elle retarde significativement la mise au rebut des nœuds capteurs, réduisant "
    "l'empreinte écotoxique liée à l'extraction des matières premières (lithium, terres rares)."
)

add_body(doc,
    "L'apiculture de précision (Precision Apiculture) représente un champ d'application "
    "emblématique pour ces technologies. L'inspection manuelle d'une ruche — nécessitant "
    "l'ouverture du couvercle — rompt brutalement l'équilibre thermique maintenu par la "
    "colonie à 34,5 °C ± 0,5 °C autour du couvain. La chute thermique à la température "
    "ambiante impose à la colonie un coût métabolique élevé (métabolisation du miel stocké "
    "pour réchauffer le nid), qui peut atteindre plusieurs dizaines de grammes de miel par "
    "intervention en période hivernale. Un système de surveillance non-intrusif basé sur des "
    "capteurs IoT LoRaWAN peut substituer partiellement ces inspections physiques, limitant "
    "le stress thermique imposé à la colonie."
)

# ── 1.2 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "1.2  Positionnement scientifique et hypothèses", level=2)

add_body(doc,
    "Ce projet implémente un Jumeau Numérique (Digital Twin) au sens strict de la définition "
    "formelle proposée par Grieves & Vickers [2] : un système tripartite composé de l'entité "
    "physique réelle (ou sa simulation biophysique dans notre cas), de son homologue virtuel "
    "calculatoire, et d'un flux de données bidirectionnel les reliant. L'originalité de notre "
    "approche réside dans la co-simulation des contraintes radio LoRaWAN (Packet Delivery Rate) "
    "et du comportement biologique de la ruche, permettant d'évaluer l'impact des pertes réseau "
    "sur la fiabilité de la détection d'événements critiques."
)

add_body(doc,
    "Deux questions scientifiques structurent l'expérimentation :"
)
add_body(doc,
    "Q1 — Dans quelle mesure un Jumeau Numérique alimenté par des données IoT peut-il prédire "
    "l'évolution d'un système biophysique réel ? Cette question cible la fidélité du modèle "
    "thermique par rapport à la référence BEEHAVE (cible 34,5 °C)."
)
add_body(doc,
    "Q2 — Les erreurs de prédiction et de détection augmentent-elles significativement en "
    "conditions extrêmes ? Cette question est intentionnellement ambiguë : elle appelle une "
    "réponse à deux niveaux orthogonaux — la stabilité des erreurs thermiques (indépendante "
    "de l'essaimage) et la dégradation de la détection d'événements discrets sous contrainte "
    "réseau (PDR dégradé)."
)

add_body(doc, "Trois hypothèses formalisées en découlent :")

add_table(doc,
    ["ID", "Énoncé", "Critère de validation"],
    [
        ["H1", "PDR = 0,90 (nominal) : détection d'essaimage très fiable.", "Précision ≥ 0,90"],
        ["H2", "PDR = 0,65 (dégradé) : chute du Rappel par lacunes LoRaWAN.", "Rappel < Rappel H1 (FN ↑)"],
        ["H3", "Chaîne MQTT → InfluxDB → Grafana stable en continu.", "Stack opérationnelle ≥ 6 h"],
    ]
)

# ── 1.3 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "1.3  Revue de littérature critique", level=2)

add_body(doc,
    "Cinq travaux fondateurs structurent le positionnement de ce projet dans l'état de l'art. "
    "Leur sélection repose sur leur pertinence directe avec les hypothèses H1/H2/H3 et la "
    "robustesse de leur méthodologie."
)

add_heading(doc, "Modélisation biophysique : BEEHAVE et Grieves & Vickers", level=3)

add_body(doc,
    "Becher et al. [1] proposent BEEHAVE, un modèle individu-centré (Agent-Based Model) "
    "simulant le comportement d'une colonie d'abeilles à l'échelle individuelle : chaque "
    "butineuse est modélisée avec ses paramètres de vol, de charge en nectar et de mortalité. "
    "Le modèle intègre un bilan énergétique complet (thermorégulation, consommation de miel, "
    "dynamique de population du couvain) validé sur des données terrain sur plusieurs saisons. "
    "Notre implémentation en est une approximation de premier ordre volontairement simplifiée : "
    "le cycle de butinage est réduit à une variation linéaire par paliers (−0,15 kg/h au départ "
    "des butineuses entre 8 h et 11 h, +0,20 kg/h au retour entre 11 h et 18 h), sans "
    "modélisation de la disponibilité florale ni variation saisonnière. Cet écart est assumé "
    "dans le cadre d'un projet académique dont l'objectif est de générer des séries temporelles "
    "suffisamment réalistes pour tester un algorithme de détection, non de reproduire BEEHAVE."
)

add_body(doc,
    "Grieves & Vickers [2] fournissent la définition formelle du Jumeau Numérique adoptée "
    "dans ce projet. Leur cadre tripartite (entité physique, jumeau virtuel, flux de données) "
    "est opérationnalisé ici par le pipeline : simulateur biophysique → MQTT → InfluxDB → "
    "Grafana + train_eval.py. L'absence de vrai capteur physique (ruche hardware) constitue "
    "la principale limite de validité externe : le jumeau est alimenté par sa propre simulation, "
    "ce qui crée une circularité contrôlée utile pour tester l'algorithme mais non transférable "
    "directement à un déploiement terrain sans validation supplémentaire."
)

add_heading(doc, "Surveillance non-intrusive et détection d'anomalies", level=3)

add_body(doc,
    "Meikle & Holst [3] démontrent expérimentalement que l'enregistrement continu du poids "
    "et de la température d'une ruche réelle permet de détecter des événements biologiques "
    "majeurs (essaimage, pic de butinage, hivernage) avec une fidélité comparable à "
    "l'observation directe. Leur travail valide l'hypothèse fondatrice de notre approche : "
    "la masse est le signal le plus informatif pour la surveillance apicole. Ils identifient "
    "également la fenêtre critique de l'essaimage — perte de 1,5 à 3 kg en moins d'une heure "
    "en pleine journée — que notre algorithme de détection par dérivée temporelle cherche à "
    "capturer."
)

add_body(doc,
    "Munir et al. [4] proposent DeepAnT, une architecture CNN non supervisée pour la "
    "détection d'anomalies dans les séries temporelles, atteignant des performances "
    "comparables aux méthodes supervisées sans nécessiter d'étiquetage. Notre approche par "
    "seuillage de dérivée (−0,03 kg/min) est délibérément plus simple et interprétable, "
    "au prix d'une moindre généralisation : elle impose de connaître a priori le régime de "
    "dynamique massique, ce qui est justifiable en simulation contrôlée mais constituerait "
    "une limite opérationnelle sur données terrain bruités."
)

add_heading(doc, "Simulation LoRaWAN : LoRaSim et passage à l'échelle", level=3)

add_body(doc,
    "Bor et al. [5] introduisent LoRaSim, le premier simulateur de réseau LoRaWAN à grande "
    "échelle, et démontrent que le PDR chute en dessous de 50 % au-delà de quelques centaines "
    "de nœuds actifs sur un même gateway, principalement en raison des collisions liées au "
    "Duty Cycle (1 % réglementaire ETSI EN 300 220). Pour un déploiement de taille modeste "
    "(< 50 nœuds), le PDR reste supérieur à 90 %, ce qui justifie notre valeur nominale "
    "PDR = 0,90 (H1). La valeur dégradée PDR = 0,65 simule un environnement forestier "
    "dense (atténuation path-loss par le feuillage) ou un rucher de taille moyenne avec "
    "congestion partielle. Un point critique soulevé par [5] concerne la nature des pertes : "
    "un canal radio réel présente des erreurs en rafales corrélées (fading de Rayleigh), "
    "alors que notre modèle applique un processus de Bernoulli indépendant par paquet — "
    "biais qui pourrait sous-estimer l'effet H2 (gaps consécutifs réels > gaps indépendants simulés)."
)

add_figure(doc, 1,
    "Architecture de la stack Docker : pipeline MQTT → Mosquitto → InfluxDB → Grafana → train_eval.py"
)

# ── 1.4 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "1.4  Contributions et plan du document", level=2)

add_body(doc,
    "Par rapport à l'état de l'art, ce projet apporte trois contributions originales dans "
    "un cadre académique reproductible (stack Docker complète, `docker compose up -d`) :"
)

add_body(doc,
    "C1 — Co-simulation biophysique + contrainte réseau : le PDR LoRaWAN est intégré "
    "directement dans la boucle de génération des données (côté publisher), permettant "
    "d'évaluer l'impact réel des lacunes de transmission sur un algorithme de détection "
    "d'événements discrets — dimension absente des travaux de Meikle & Holst [3] qui "
    "opèrent sur des séries temporelles continues."
)

add_body(doc,
    "C2 — Résultat contre-intuitif sur l'orthogonalité des erreurs : les erreurs thermiques "
    "(MAE = 0,635 °C) restent stables entre régime nominal et extrême (ΔMAE < 0,01 °C), "
    "tandis que les erreurs de détection augmentent sous contrainte réseau (Rappel : 1,00 → 0,94). "
    "Ces deux dimensions d'erreur sont orthogonales, ce qui constitue le résultat scientifique "
    "central de ce dossier."
)

add_body(doc,
    "C3 — Infrastructure de recherche ouverte : la stack conteneurisée (Mosquitto, InfluxDB 2.7, "
    "Grafana, Python 3.x) est intégralement reproductible sans configuration manuelle, "
    "offrant une base de travail réutilisable pour des expérimentations futures sur des "
    "données capteurs réels."
)

add_body(doc,
    "Le reste du dossier est organisé comme suit : le Chapitre 2 analyse la couche physique "
    "LoRa (modulation CSS, budget de liaison, ADR, modèle PDR). Le Chapitre 3 couvre le "
    "protocole MAC LoRaWAN, la sécurité AES-128, la scalabilité et le transport MQTT. "
    "Le Chapitre 4 présente les résultats de performance (thermorégulation, détection H1/H2) "
    "et les limites de conception. Le Chapitre 5 ouvre sur les perspectives R&D et conclut."
)

# ---------------------------------------------------------------------------
# Sauvegarde
# ---------------------------------------------------------------------------
out_path = "Chapitre_1_LoRaWAN.docx"
doc.save(out_path)
print(f"✓ Document généré : {out_path}")
