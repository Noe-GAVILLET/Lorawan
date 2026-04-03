"""
Génère Chapitre_5_LoRaWAN.docx  (~3 pages)
Chapitre 5 — Perspectives R&D et Conclusion
Projet ETRS012 — Jumeau Numérique Scientifique de Ruche IoT
Université Savoie Mont Blanc — Avril 2026
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()


def add_formula(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()

# ── Titre du chapitre ───────────────────────────────────────────────────────
add_heading(doc, "Chapitre 5 — Perspectives R&D et Conclusion", level=1)

add_body(doc,
    "Ce dernier chapitre prolonge les résultats du Chapitre 4 vers leurs implications "
    "pratiques et scientifiques. Il identifie les voies de maturation technique du "
    "système pour un déploiement terrain réel, les améliorations algorithmiques "
    "prioritaires, et les enjeux éthiques et environnementaux qui conditionnent "
    "l'acceptabilité sociale du projet. La conclusion synthétise les réponses aux "
    "deux questions scientifiques initiales et positionne la contribution du projet "
    "dans l'état de l'art."
)

# ── 5.1 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "5.1  Vers un déploiement LoRaWAN réel", level=2)

add_body(doc,
    "L'architecture logicielle du projet a été conçue dès l'origine pour faciliter "
    "la migration vers un déploiement terrain réel, sans modification du pipeline aval. "
    "La substitution du simulateur biophysique par un vrai nœud LoRa physique se résume "
    "à un unique point d'entrée dans le système : le remplacement de random_data_publisher.py "
    "par un nœud matériel (microcontrôleur + module radio LoRa) transmettant vers une "
    "gateway LoRaWAN connectée à The Things Network (TTN) ou un Network Server privé "
    "(ChirpStack). TTN expose un broker MQTT natif sur lequel mqtt_to_influx.py peut "
    "s'abonner sans modification, en changeant uniquement l'URL du broker et le topic "
    "(`v3/{app_id}/devices/{dev_id}/up` au lieu de `hive/ruche-01/telemetry`)."
)

add_body(doc,
    "Le tableau suivant estime l'effort de migration par composant :"
)

add_table(doc,
    ["Composant", "Statut actuel", "Modification requise pour déploiement réel", "Effort"],
    [
        ["random_data_publisher.py", "Simulateur logiciel",      "Remplacé par nœud LoRa physique (ex. Adafruit Feather M0 LoRa + capteur HX711 + DS18B20)", "Hardware + firmware C/Arduino"],
        ["mqtt_to_influx.py",        "Opérationnel",             "Changer MQTT_HOST + topic TTN (format JSON TTN v3)",                                         "< 1h développement"],
        ["InfluxDB + Grafana",       "Opérationnel",             "Aucune modification — API Flux identique",                                                    "Zéro"],
        ["train_eval.py",            "Opérationnel",             "Recalibration du seuil −0,03 kg/min sur données réelles terrain",                             "Expérimentation terrain"],
        ["Sécurité MQTT",            "En clair (lab)",           "Activer TLS (MQTTS :8883) + certificats TTN",                                                 "< 2h configuration"],
    ]
)

add_body(doc,
    "La recalibration du seuil de détection (train_eval.py) constitue l'effort le plus "
    "substantiel : les dérivées réelles de masse d'une ruche terrain incluent du bruit "
    "mécanique (vibrations, pluie, vent), des dérives de capteur thermiques et des "
    "pics d'activité non linéaires absents du simulateur. Une campagne de collecte sur "
    "une saison complète (mars–septembre) avec N ≥ 3 ruches instrumentées serait "
    "nécessaire pour établir des distributions de dérivées par régime et recalibrer "
    "le seuil par analyse ROC (Receiver Operating Characteristic) plutôt que par "
    "fixation manuelle — éliminant ainsi la limite L4 identifiée en § 4.6."
)

# ── 5.2 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "5.2  Perspectives algorithmiques", level=2)

add_heading(doc, "Du seuillage vers l'apprentissage automatique", level=3)

add_body(doc,
    "L'algorithme de détection par dérivée temporelle (§ 4.3) est intentionnellement "
    "simple et interprétable. Trois axes d'amélioration algorithmique sont identifiés "
    "par ordre de complexité croissante :"
)

add_body(doc,
    "Axe A — Seuil adaptatif par régression quantile : plutôt qu'un seuil fixe "
    "−0,03 kg/min, calibrer le seuil dynamiquement sur un percentile bas (P5) de la "
    "distribution de dérivées observées en régime nominal (fenêtre glissante de 24 h). "
    "Cela rendrait le détecteur robuste aux dérives lentes de masse (hivernage progressif, "
    "évaporation) sans nécessiter de recalibration manuelle. Implémentable en < 50 lignes "
    "Python avec scipy.stats.scoreatpercentile, sans entraînement."
)

add_body(doc,
    "Axe B — Détection par LSTM (Long Short-Term Memory) : les architectures LSTM "
    "sont particulièrement adaptées à la détection d'anomalies dans les séries temporelles "
    "multivariées (masse + température + heure du jour). Entraîné sur des périodes nominales "
    "uniquement (approche one-class), un LSTM prédit la masse attendue M̂(t+1) et signale "
    "une anomalie lorsque |M(t+1) − M̂(t+1)| dépasse plusieurs écarts-types. Cette approche "
    "— implémentable avec TensorFlow Lite pour embarquement sur microcontrôleur ARM Cortex-M4 "
    "— offre une meilleure généralisation que le seuillage mais au prix d'une explicabilité "
    "réduite et d'un besoin de données d'entraînement labellisées."
)

add_body(doc,
    "Axe C — DeepAnT (Munir et al. [4]) : l'architecture CNN non supervisée de référence "
    "en détection d'anomalies sur séries temporelles compressées. Son avantage est l'absence "
    "de nécessité d'exemples d'essaimage labellisés — particulièrement utile pour les "
    "espèces ou les régions où l'essaimage est rare. Son application directe à notre "
    "problème nécessiterait cependant de gérer les données manquantes (PDR < 1) en amont, "
    "par des techniques d'imputation plus sophistiquées que l'interpolation linéaire "
    "(ex. imputation par modèle gaussien ou par GAN conditionnel)."
)

add_heading(doc, "Robustesse statistique : vers N = 30 répétitions", level=3)

add_body(doc,
    "La limite L1 (§ 4.6) — absence d'intervalle de confiance — est la plus simple à "
    "corriger sans modification algorithmique. Il suffit de :"
)

add_formula(doc,
    "for seed in range(30):\n"
    "    random.seed(seed)\n"
    "    run_simulation()  # génère dataset H1 ou H2\n"
    "    compute_metrics() # calcule Précision, Rappel, F1\n"
    "→ Moyenne ± 1,96 σ/√30  (IC 95 %)"
)

add_body(doc,
    "Sur N = 30 runs avec graine fixée, la distribution du Rappel H2 autour de 0,94 "
    "permettrait de conclure : 'Le Rappel en conditions H2 est μ ± σ avec p < 0,05 "
    "par rapport à H1 (test de Wilcoxon sur données appariées)' — résultat publiable "
    "dans un atelier IoT ou un poster de conférence nationale."
)

add_heading(doc, "Détection acoustique : chant des reines", level=3)

add_body(doc,
    "Une limite fondamentale du système actuel est l'absence de modalité acoustique. "
    "Le chant des reines — émis par les reines vierges dans les jours précédant "
    "l'essaimage (fréquence de 'tooting' : 350–500 Hz, 'quacking' : 300–400 Hz) — "
    "constitue un signal précurseur observable jusqu'à 48 h avant l'essaimage, bien "
    "avant toute variation de masse détectable. L'ajout d'un microphone MEMS (ex. "
    "ICS43434, ~1 €) et d'une FFT embarquée sur le nœud LoRa permettrait une fusion "
    "multimodale masse + acoustique, réduisant la fenêtre de détection de quelques "
    "heures à quelques jours et éliminant la dépendance critique au τ_pub et au PDR "
    "pour les alertes précoces."
)

# ── 5.3 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "5.3  Éthique et impact environnemental", level=2)

add_body(doc,
    "Le déploiement d'un système IoT en milieu naturel, aussi bien intentionné soit-il, "
    "génère une empreinte matérielle et numérique qui doit être évaluée avec honnêteté "
    "dans une démarche de numérique responsable."
)

add_heading(doc, "Analyse du Cycle de Vie (ACV) des composants", level=3)

add_body(doc,
    "Un nœud LoRa typique (microcontrôleur STM32 + module SX1276 + cellule HX711 + "
    "batterie LiPo 3,7 V / 3 Ah) représente environ 15 à 25 g de matière électronique "
    "contenant du lithium, du cuivre, des terres rares (europium, terbium pour les "
    "LED de status) et des polymères halogénés dans les PCB. La phase de fabrication "
    "(principalement en Asie du Sud-Est) génère l'essentiel de l'empreinte carbone "
    "du composant — estimée à 3–8 kg CO₂e par nœud selon les analyses ACV disponibles "
    "pour des produits électroniques similaires."
)

add_body(doc,
    "C'est ici que LoRaWAN révèle son avantage environnemental décisif par rapport "
    "aux alternatives : une durée de vie de 3 à 5 ans sur batterie (contre 3 à 7 jours "
    "pour un module Wi-Fi ou BLE en transmission continue) amortit l'empreinte de "
    "fabrication sur une période longue, réduisant le flux de déchets électroniques "
    "générés par ruche instrumentée d'un facteur 200 à 600 par rapport à un déploiement "
    "4G/LTE équivalent. En appliquant le principe de sobriété numérique — émettre "
    "le minimum de données nécessaires à la décision — LoRaWAN s'impose comme le "
    "protocole le plus aligné avec les objectifs de développement durable appliqués "
    "à l'agriculture connectée."
)

add_heading(doc, "Gouvernance des données agricoles", level=3)

add_body(doc,
    "Les données produites par un réseau de ruches connectées — masse, température, "
    "localisation — possèdent une dimension concurrentielle et sécuritaire souvent "
    "sous-estimée. La connaissance précise du moment optimal de récolte (pic de masse "
    "juste avant essaimage) ou de la localisation GPS d'un rucher à haute valeur "
    "productive expose l'apiculteur à un risque de vol ciblé. Le chiffrement "
    "AES-128 natif de LoRaWAN (§ 3.3) protège les données en transit, mais la "
    "sécurisation de bout en bout jusqu'à l'application cliente nécessite une "
    "implémentation TLS correcte sur l'ensemble de la chaîne (MQTTS, HTTPS pour "
    "les API Grafana/InfluxDB)."
)

add_body(doc,
    "Dans une perspective de science ouverte et de coopération avec les organismes "
    "de recherche publics (CNRS, INRAE), une architecture de données anonymisées "
    "par région géographique (agrégation à l'échelle de la commune) permettrait de "
    "partager les courbes de poids et de butinage avec les équipes travaillant sur "
    "l'impact du dérèglement climatique sur les floraisons et la phénologie des "
    "pollinisateurs — sans jamais exposer la localisation précise des ruchers "
    "individuels. Ce modèle de 'données ouvertes agrégées' est techniquement "
    "réalisable via une étape de généralisation géographique avant l'export CSV, "
    "en supprimant les coordonnées GPS brutes et en ne conservant qu'un identifiant "
    "de zone (code INSEE ou maille 10 × 10 km²)."
)

# ── 5.4 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "5.4  Conclusion générale", level=2)

add_body(doc,
    "Ce dossier a présenté la conception, l'implémentation et l'évaluation rigoureuse "
    "d'un Jumeau Numérique de ruche connectée, articulant un modèle biophysique "
    "BEEHAVE-like, une simulation des contraintes radio LoRaWAN, et un algorithme "
    "de détection d'essaimage par dérivée temporelle. L'ensemble de la stack — "
    "Mosquitto, InfluxDB 2.7, Grafana, Python 3.x — est intégralement reproductible "
    "via `docker compose up -d`, sans configuration manuelle."
)

add_body(doc,
    "Les réponses aux deux questions scientifiques initiales sont les suivantes :"
)

add_table(doc,
    ["Question", "Réponse synthétique", "Résultat quantitatif"],
    [
        ["Q1 — Un Jumeau Numérique peut-il prédire l'évolution du système physique ?",
         "Oui, avec une fidélité thermique bien en dessous du critère d'acceptance, et une détection d'événements quasi-parfaite en conditions nominales.",
         "MAE = 0,635 °C (< 1,0 °C requis) ; F1 = 1,00 (H1) ; F1 = 0,97 (H2)"],
        ["Q2 — Les erreurs augmentent-elles en conditions extrêmes ?",
         "Partiellement — les erreurs thermiques restent stables (découplage masse/thermorégulation), mais les erreurs de détection augmentent sous contrainte réseau.",
         "ΔMAE thermique < 0,01 °C (stable) ; Rappel : 1,00 → 0,94 (dégradé sous PDR=0,65)"],
    ]
)

add_body(doc,
    "Le résultat scientifique original de ce projet est l'orthogonalité démontrée des "
    "deux dimensions d'erreur : un essaimage ne dégrade pas la précision thermique "
    "du Jumeau Numérique, mais les pertes radio LoRaWAN dégradent bien la capacité "
    "à détecter cet événement. Ces deux facettes — précision du modèle biophysique "
    "et robustesse de la détection sous contrainte réseau — sont indépendantes et "
    "doivent être évaluées séparément pour tout système IoT de surveillance à "
    "événements discrets."
)

add_body(doc,
    "Trois perspectives de maturation ont été identifiées comme prioritaires pour "
    "franchir le seuil de la recherche appliquée au déploiement opérationnel : "
    "(i) la migration vers un nœud LoRa physique avec recalibration terrain du "
    "seuil de détection ; (ii) l'établissement d'intervalles de confiance sur "
    "N ≥ 30 répétitions pour valider la robustesse statistique des métriques ; "
    "(iii) l'intégration d'une modalité acoustique (chant des reines) pour "
    "fournir une alerte précoce indépendante des contraintes de Duty Cycle LoRaWAN."
)

add_body(doc,
    "Au-delà des aspects techniques, ce projet illustre une tension fondamentale "
    "du numérique responsable appliqué à l'agriculture : la technologie doit rester "
    "au service du bien-être animal et de l'autonomie de l'apiculteur, et non "
    "l'inverse. Un système d'alerte qui génère des fausses alertes excessives, "
    "consomme les batteries en quelques semaines, ou expose les données de l'apiculteur "
    "à des tiers non consentis n'est pas éthiquement justifiable, quelle que soit "
    "sa sophistication algorithmique. La sobriété de conception — une dérivée "
    "temporelle, un seuil, une batterie qui dure 3 ans — est une réponse "
    "d'ingénierie, mais aussi un choix éthique."
)

# ── Références bibliographiques ─────────────────────────────────────────────
add_heading(doc, "Références bibliographiques", level=1)

refs = [
    "[1] Becher, M. A., Grimm, V., Thorbek, P., Horn, J., Kennedy, P. J., & Osborne, J. L. (2014). "
    "BEEHAVE: a systems model of honeybee colony dynamics and foraging to explore multifactorial "
    "causes of colony failure. Journal of Applied Ecology, 51(2), 470–482. "
    "https://doi.org/10.1111/1365-2664.12222",

    "[2] Grieves, M., & Vickers, J. (2017). Digital Twin: Mitigating Unpredictable, Undesirable "
    "Emergent Behavior in Complex Systems. In Transdisciplinary Perspectives on Complex Systems "
    "(pp. 85–113). Springer. https://doi.org/10.1007/978-3-319-38756-7_4",

    "[3] Meikle, W. G., & Holst, N. (2015). Application of continuous monitoring of honeybee "
    "colonies. Apidologie, 46(1), 10–22. https://doi.org/10.1007/s13592-014-0298-x",

    "[4] Munir, M., Siddiqui, S. A., Dengel, A., & Ahmed, S. (2019). DeepAnT: A deep learning "
    "approach for unsupervised anomaly detection in time series. IEEE Access, 7, 1991–2005. "
    "https://doi.org/10.1109/ACCESS.2018.2886457",

    "[5] Bor, M. C., Roedig, U., Voigt, T., & Alonso, J. M. (2016). Do LoRa Low-Power Wide-Area "
    "Networks Scale? In Proceedings of the 19th ACM International Conference on Modeling, Analysis "
    "and Simulation of Wireless and Mobile Systems (MSWiM '16), pp. 59–67. ACM. "
    "https://doi.org/10.1145/2988287.2989163",
]

for ref in refs:
    p = doc.add_paragraph(ref, style="Normal")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(4)

# ---------------------------------------------------------------------------
# Sauvegarde
# ---------------------------------------------------------------------------
out_path = "Chapitre_5_LoRaWAN.docx"
doc.save(out_path)
print(f"✓ Document généré : {out_path}")
