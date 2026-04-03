"""
Génère Page_de_garde_LoRaWAN.docx
Page de garde — Dossier Technique R&D
Projet ETRS012 — Jumeau Numérique Scientifique de Ruche IoT
Université Savoie Mont Blanc — Avril 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers bas niveau
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Couleur de fond d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 6)))
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def para_space(doc, size_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.font.size = Pt(size_pt)


def set_page_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


# Couleurs du projet
NAVY      = "1B3A5C"   # bleu marine profond
GOLD      = "D4A017"   # or apicole
LIGHTGRAY = "F2F4F7"   # gris très clair
WHITE     = "FFFFFF"
DARKGRAY  = "4A4A4A"


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()
set_page_margins(doc, top=0, bottom=0, left=0, right=0)

# ── Bande supérieure marine pleine largeur ───────────────────────────────────
top_table = doc.add_table(rows=1, cols=1)
top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
top_cell = top_table.rows[0].cells[0]
set_cell_bg(top_cell, NAVY)
top_cell.width = Cm(21)

p = top_cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n")
run.font.size = Pt(28)

# Logo texte université
p2 = top_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("UNIVERSITÉ SAVOIE MONT BLANC")
r.font.bold  = True
r.font.size  = Pt(13)
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
r.font.name  = "Calibri"

p3 = top_cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run("Master 2 Réseaux & Télécommunications  —  ETRS012 Intelligence Ambiante")
r.font.size  = Pt(10)
r.font.color.rgb = RGBColor(0xCC, 0xD6, 0xE8)
r.font.name  = "Calibri"
p3.paragraph_format.space_after = Pt(18)

doc.add_paragraph()  # espace

# ── Ligne dorée décorative ────────────────────────────────────────────────────
gold_table = doc.add_table(rows=1, cols=1)
gold_table.alignment = WD_TABLE_ALIGNMENT.CENTER
gc = gold_table.rows[0].cells[0]
set_cell_bg(gc, GOLD)
p_g = gc.paragraphs[0]
p_g.paragraph_format.space_before = Pt(0)
p_g.paragraph_format.space_after  = Pt(0)
run_g = p_g.add_run(" ")
run_g.font.size = Pt(4)

doc.add_paragraph()

# ── Pictogramme ruche (texte stylisé) ────────────────────────────────────────
p_icon = doc.add_paragraph()
p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_icon = p_icon.add_run("🐝")
r_icon.font.size = Pt(52)
p_icon.paragraph_format.space_before = Pt(10)
p_icon.paragraph_format.space_after  = Pt(6)

# ── Titre principal ───────────────────────────────────────────────────────────
p_titre = doc.add_paragraph()
p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p_titre.add_run("Jumeau Numérique")
r1.font.size  = Pt(32)
r1.font.bold  = True
r1.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
r1.font.name  = "Calibri"
p_titre.paragraph_format.space_before = Pt(0)
p_titre.paragraph_format.space_after  = Pt(0)

p_titre2 = doc.add_paragraph()
p_titre2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_titre2.add_run("d'une Ruche Connectée")
r2.font.size  = Pt(32)
r2.font.bold  = True
r2.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
r2.font.name  = "Calibri"
p_titre2.paragraph_format.space_before = Pt(0)
p_titre2.paragraph_format.space_after  = Pt(8)

# Sous-titre
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run("Analyse R&D d'un système IoT LoRaWAN")
r_sub.font.size  = Pt(16)
r_sub.font.bold  = False
r_sub.font.italic = True
r_sub.font.color.rgb = RGBColor(0xD4, 0xA0, 0x17)
r_sub.font.name  = "Calibri"
p_sub.paragraph_format.space_after = Pt(20)

# ── Ligne dorée fine ──────────────────────────────────────────────────────────
gold_table2 = doc.add_table(rows=1, cols=1)
gold_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
gc2 = gold_table2.rows[0].cells[0]
set_cell_bg(gc2, GOLD)
p_g2 = gc2.paragraphs[0]
p_g2.paragraph_format.space_before = Pt(0)
p_g2.paragraph_format.space_after  = Pt(0)
run_g2 = p_g2.add_run(" ")
run_g2.font.size = Pt(3)

doc.add_paragraph()

# ── Bloc résumé (carte grise) ─────────────────────────────────────────────────
card = doc.add_table(rows=1, cols=1)
card.alignment = WD_TABLE_ALIGNMENT.CENTER
card_cell = card.rows[0].cells[0]
set_cell_bg(card_cell, LIGHTGRAY)
set_cell_borders(card_cell,
    top    ={"val": "single", "sz": 12, "color": GOLD},
    bottom ={"val": "single", "sz": 12, "color": GOLD},
    left   ={"val": "single", "sz": 12, "color": NAVY},
    right  ={"val": "single", "sz": 12, "color": NAVY},
)

def card_line(cell, label, value, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    rl = p.add_run(f"{label}  ")
    rl.font.bold  = True
    rl.font.size  = Pt(11)
    rl.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    rl.font.name  = "Calibri"
    rv = p.add_run(value)
    rv.font.size  = Pt(11)
    rv.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    rv.font.name  = "Calibri"

card_line(card_cell, "Auteur :",       "Noé Gavillet", first=True)
card_line(card_cell, "Module :",       "ETRS012 — Intelligence Ambiante")
card_line(card_cell, "Encadrement :",  "Université Savoie Mont Blanc")
card_line(card_cell, "Date :",         "Avril 2026")
card_line(card_cell, "Version :",      "1.0 — Dossier Technique R&D (20 pages)")

p_pad = card_cell.add_paragraph()
p_pad.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# ── Mots-clés ─────────────────────────────────────────────────────────────────
p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.CENTER
kw_label = p_kw.add_run("Mots-clés : ")
kw_label.font.bold  = True
kw_label.font.size  = Pt(10)
kw_label.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
kw_val = p_kw.add_run(
    "LoRaWAN · Digital Twin · LPWAN · IoT · Apiculture de précision · "
    "Détection d'essaimage · BEEHAVE · Séries temporelles · AES-128 · MQTT · InfluxDB"
)
kw_val.font.size  = Pt(10)
kw_val.font.italic = True
kw_val.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
p_kw.paragraph_format.space_after = Pt(0)

# ── Bande inférieure marine ────────────────────────────────────────────────────
doc.add_paragraph()

bot_table = doc.add_table(rows=1, cols=3)
bot_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Colonne gauche
bc_left = bot_table.rows[0].cells[0]
set_cell_bg(bc_left, NAVY)
p_bl = bc_left.paragraphs[0]
p_bl.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_bl.paragraph_format.space_before = Pt(8)
p_bl.paragraph_format.space_after  = Pt(8)
r_bl = p_bl.add_run("  ETRS012 — 2025/2026")
r_bl.font.size  = Pt(9)
r_bl.font.color.rgb = RGBColor(0xCC, 0xD6, 0xE8)
r_bl.font.name  = "Calibri"

# Colonne centre (ligne dorée verticale symbolique)
bc_mid = bot_table.rows[0].cells[1]
set_cell_bg(bc_mid, GOLD)
p_bm = bc_mid.paragraphs[0]
p_bm.paragraph_format.space_before = Pt(0)
p_bm.paragraph_format.space_after  = Pt(0)
r_bm = p_bm.add_run(" ")
r_bm.font.size = Pt(1)

# Colonne droite
bc_right = bot_table.rows[0].cells[2]
set_cell_bg(bc_right, NAVY)
p_br = bc_right.paragraphs[0]
p_br.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_br.paragraph_format.space_before = Pt(8)
p_br.paragraph_format.space_after  = Pt(8)
r_br = p_br.add_run("Confidentiel — Usage académique  ")
r_br.font.size  = Pt(9)
r_br.font.color.rgb = RGBColor(0xCC, 0xD6, 0xE8)
r_br.font.name  = "Calibri"

# Largeurs colonnes bande inférieure
bot_table.columns[0].width = Cm(9)
bot_table.columns[1].width = Cm(1)
bot_table.columns[2].width = Cm(9)

# ---------------------------------------------------------------------------
# Sauvegarde
# ---------------------------------------------------------------------------
out_path = "Page_de_garde_LoRaWAN.docx"
doc.save(out_path)
print(f"✓ Page de garde générée : {out_path}")
