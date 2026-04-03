"""
Génère Chapitre_3_LoRaWAN.docx  (~5 pages)
Chapitre 3 — Protocole MAC, Sécurité et Passage à l'échelle
Projet ETRS012 — Jumeau Numérique Scientifique de Ruche IoT
Université Savoie Mont Blanc — Avril 2026
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_figure(doc, number, description):
    p = doc.add_paragraph()
    run = p.add_run(f"[FIGURE {number} : {description}]")
    run.bold = True
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
    doc.add_paragraph()


def add_formula(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()

# ── Titre du chapitre ───────────────────────────────────────────────────────
add_heading(doc, "Chapitre 3 — Protocole MAC, Sécurité et Passage à l'Échelle", level=1)

add_body(doc,
    "Si le Chapitre 2 a analysé la couche physique LoRa (CSS, budget de liaison, SF), "
    "ce chapitre monte d'un niveau d'abstraction pour examiner le protocole d'accès au "
    "médium (MAC) LoRaWAN, son architecture de sécurité cryptographique, les contraintes "
    "réglementaires de Duty Cycle, la scalabilité du réseau selon les résultats de LoRaSim, "
    "et la couche de transport MQTT utilisée dans la stack expérimentale. Une comparaison "
    "critique des technologies LPWAN concurrentes clôt le chapitre."
)

# ── 3.1 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.1  Structure MAC LoRaWAN : classes, trames et fenêtres de réception", level=2)

add_body(doc,
    "LoRaWAN définit trois classes d'opération pour les nœuds terminaux (End Devices), "
    "chacune répondant à un compromis différent entre latence descendante et consommation "
    "énergétique :"
)

add_table(doc,
    ["Classe", "Comportement réception", "Latence downlink", "Consommation"],
    [
        ["A (All)", "RX1 + RX2 après chaque uplink uniquement", "Variable (après prochain uplink)", "Minimale — nœud en veille profonde"],
        ["B (Beacon)", "Créneaux RX périodiques synchronisés sur beacon gateway", "Prévisible (< 128 s)", "Intermédiaire"],
        ["C (Continuous)", "Réception quasi-continue, sauf pendant l'émission", "Quasi-nulle (< 1 s)", "Maximale — batterie inadaptée"],
    ]
)

add_body(doc,
    "Pour un capteur de ruche sur batterie, la Classe A est impérative. Le cycle d'une "
    "trame Classe A se déroule ainsi : le nœud émet un uplink (UL), puis ouvre deux "
    "fenêtres de réception successives. RX1 s'ouvre exactement RECEIVE_DELAY1 secondes "
    "après la fin de l'uplink (valeur par défaut : 1 s) sur la même fréquence et le même "
    "SF. RX2 s'ouvre RECEIVE_DELAY2 secondes après l'UL (défaut : 2 s) sur un canal de "
    "secours fixe à 869,525 MHz, SF12, pour maximiser la portée du downlink. Si le "
    "Network Server n'a pas de commande à envoyer, les deux fenêtres se ferment sans "
    "activité radio. Ce mécanisme permet au nœud de dormir l'essentiel du temps, "
    "consommant < 1 µA en veille contre ~45 mA en émission."
)

add_body(doc,
    "La structure d'une trame LoRaWAN (PHYPayload) est la suivante :"
)

add_formula(doc,
    "PHYPayload = MHDR (1B) | MACPayload | MIC (4B)\n"
    "MACPayload  = FHDR (7–22B) | FPort (1B, optionnel) | FRMPayload (données chiffrées)\n"
    "FHDR        = DevAddr (4B) | FCtrl (1B) | FCnt (2B) | FOpts (0–15B)"
)

add_body(doc,
    "Le champ FCnt (Frame Counter) est critique pour la sécurité : il prévient les "
    "attaques par rejeu en invalidant tout paquet dont le compteur n'est pas strictement "
    "supérieur au dernier compteur reçu. Le MIC (Message Integrity Code) de 4 octets "
    "est calculé par AES-128 CMAC sur l'intégralité de la trame, garantissant l'intégrité "
    "et l'authenticité. Le FRMPayload est chiffré séparément par AES-128 CTR mode, "
    "assurant la confidentialité des données capteur. Ces deux opérations cryptographiques "
    "utilisent des clés distinctes, détaillées en section 3.3."
)

add_figure(doc, 3,
    "Structure d'une trame LoRaWAN (PHYPayload) : décomposition des champs MHDR, FHDR "
    "(DevAddr, FCnt, FOpts), FPort, FRMPayload chiffré AES-128 CTR et MIC AES-128 CMAC (4B)"
)

# ── 3.2 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.2  Contrainte Duty Cycle ISM 868 MHz et débit utile", level=2)

add_body(doc,
    "La bande ISM 868 MHz utilisée en Europe est soumise à la réglementation ETSI "
    "EN 300 220-2 V3.2.1, qui impose un Duty Cycle (DC) maximum par sous-bande. "
    "Pour la sous-bande principale 868,0–868,6 MHz (canaux obligatoires LoRaWAN), "
    "le DC est limité à 1 % — soit 36 secondes d'émission autorisée par heure. "
    "Pour la sous-bande 869,4–869,65 MHz (canal RX2 secours), le DC monte à 10 %, "
    "mais avec une puissance limitée à 27 dBm EIRP."
)

add_body(doc,
    "Le débit utile maximal par nœud en uplink est donc borné par :"
)

add_formula(doc,
    "Débit_max = DC × (Payload_max / ToA) = 0,01 × (51 B / ToA_SF)\n"
    "→ SF7  : 0,01 × (20 B / 51 ms)  ≈ 3,9 msg/min  (intervalle min : 5,1 s)\n"
    "→ SF12 : 0,01 × (20 B / 1318 ms) ≈ 0,45 msg/min (intervalle min : 132 s)"
)

add_body(doc,
    "Dans notre expérimentation, le paramètre PUBLISH_INTERVAL_SECONDS est fixé à 5 s "
    "pour le dataset H1 — à la limite basse du DC réglementaire pour SF7. Cette valeur "
    "n'est pas réaliste pour un vrai déploiement LoRaWAN en SF élevé : elle a été choisie "
    "pour générer un volume de données suffisant (5 216 points sur 8 h) pour les calculs "
    "de métriques. En production, l'intervalle réaliste de 60 s utilisé pour le dataset H2 "
    "est bien plus représentatif d'un nœud LoRaWAN conforme ETSI en SF9–SF10."
)

add_body(doc,
    "Cette contrainte réglementaire a également un impact sur la conception du système "
    "d'alerte : un nœud LoRaWAN ne peut pas augmenter dynamiquement sa cadence d'envoi "
    "lors d'un essaimage détecté localement pour accélérer la remontée d'information — "
    "il reste borné par le DC. Certaines implémentations avancées contournent ce problème "
    "par un mécanisme de 'burst reporting' sur la sous-bande 10 % (869,5 MHz), mais cette "
    "fonctionnalité n'est pas standardisée dans LoRaWAN v1.0.x."
)

# ── 3.3 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.3  Architecture de sécurité LoRaWAN", level=2)

add_heading(doc, "Procédures d'activation : OTAA vs ABP", level=3)

add_body(doc,
    "LoRaWAN définit deux mécanismes d'activation des nœuds terminaux :"
)

add_body(doc,
    "Over-The-Air Activation (OTAA) : le nœud émet une trame Join-Request contenant son "
    "DevEUI (identifiant unique 64 bits, analogue à une adresse MAC), son AppEUI (identifiant "
    "de l'application) et un DevNonce aléatoire 16 bits anti-rejeu. Le Join Server valide "
    "la requête via un MIC calculé avec l'AppKey (clé racine 128 bits, pré-provisionnée). "
    "En cas de succès, il génère deux clés de session dérivées par KDF (Key Derivation "
    "Function) : NwkSKey (Network Session Key, pour le MIC MAC) et AppSKey (Application "
    "Session Key, pour le chiffrement FRMPayload). Ces clés sont renouvelées à chaque "
    "Join-Request, ce qui offre une forward secrecy partielle."
)

add_body(doc,
    "Activation By Personalization (ABP) : les clés NwkSKey, AppSKey et le DevAddr sont "
    "provisionnées statiquement en mémoire flash du nœud. Bien que plus simple à déployer, "
    "ABP présente un risque de sécurité majeur : l'absence de renouvellement de clés et "
    "la vulnérabilité au dépassement du compteur FCnt (en cas de reset du nœud, les "
    "paquets sont rejetés par le NS car FCnt repart de 0). Pour un capteur de ruche en "
    "déploiement pluriannuel, OTAA est fortement recommandé."
)

add_heading(doc, "Chiffrement et intégrité : AES-128 double couche", level=3)

add_body(doc,
    "La sécurité LoRaWAN repose sur AES-128 appliqué en deux opérations distinctes "
    "sur chaque trame uplink :"
)

add_body(doc,
    "1. Chiffrement du FRMPayload (confidentialité) : le payload applicatif est chiffré "
    "par AES-128 en mode CTR (Counter mode) avec l'AppSKey comme clé. Le vecteur "
    "d'initialisation incorpore DevAddr, FCnt et la direction (uplink/downlink), "
    "garantissant l'unicité du keystream par message. Le mode CTR transforme AES en "
    "chiffrement de flux, permettant de traiter des payloads de longueur variable "
    "sans padding. Pour un payload de 20 octets (poids + température + timestamp), "
    "un seul bloc AES-128 suffit."
)

add_body(doc,
    "2. Calcul du MIC (intégrité et authenticité) : le Message Integrity Code de 4 octets "
    "est calculé par AES-128 CMAC (Cipher-based MAC, RFC 4493) sur la trame complète "
    "(MHDR + FHDR + FPort + FRMPayload chiffré) avec la NwkSKey. Le résultat de 16 octets "
    "est tronqué aux 4 premiers octets. Le Network Server recalcule le MIC à la réception "
    "et rejette toute trame dont le MIC ne correspond pas. Cette troncature à 4 octets "
    "implique une probabilité de collision de 2^{-32} ≈ 2,3×10^{-10} par paquet — "
    "négligeable pour un nœud émettant 1 paquet par minute."
)

add_heading(doc, "Pertinence pour les données agricoles sensibles", level=3)

add_body(doc,
    "Dans le contexte de l'apiculture connectée, les données IoT (masse, température, "
    "localisation GPS de la ruche) ont une valeur commerciale et sécuritaire non négligeable : "
    "elles permettent d'identifier précisément quand une ruche est pleine (pic de masse "
    "juste avant l'essaimage) et donc vulnérable au vol. Le chiffrement AES-128 embarqué "
    "dans LoRaWAN protège ces données en transit sur le canal radio. Toutefois, la sécurité "
    "de bout en bout (nœud → application cliente de l'apiculteur) dépend de la sécurisation "
    "du Join Server, du Application Server et des API REST en aval — éléments hors périmètre "
    "de notre stack expérimentale. Dans notre implémentation MQTT, les messages transitent "
    "en clair sur le broker Mosquitto local (QoS 0, sans TLS) : acceptable en lab isolé, "
    "inacceptable en déploiement production."
)

add_body(doc,
    "Le document éthique du projet [docs/ethique.md] soulève également la question de "
    "la gouvernance des données apicoles : une anonymisation régionale des poids collectés "
    "pourrait permettre de partager ces données avec des organismes de recherche (CNRS, INRAE) "
    "pour le suivi de l'impact du dérèglement climatique sur les floraisons, tout en "
    "protégeant l'identité et la localisation des apiculteurs participants."
)

# ── 3.4 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.4  Scalabilité et résultats LoRaSim", level=2)

add_body(doc,
    "La question de la scalabilité de LoRaWAN est traitée quantitativement par Bor et al. [5] "
    "via LoRaSim, le premier simulateur de réseau LoRaWAN à grande échelle publié dans la "
    "littérature académique. Les auteurs modélisent un réseau de N nœuds (N variant de 1 à "
    "5 000) autour d'un unique gateway, avec un intervalle de publication fixe et différents "
    "SF. Leurs résultats mettent en évidence trois régimes opérationnels distincts :"
)

add_table(doc,
    ["Nb nœuds (même GW)", "PDR observé (SF mixte)", "Cause principale de dégradation"],
    [
        ["< 50",          "> 95 %",    "Collisions quasi-nulles, DC non contraignant"],
        ["50 – 500",      "50 – 90 %", "Collisions en mode Aloha (même SF, même canal)"],
        ["> 500",         "< 50 %",    "Saturation DC gateway, collisions en cascade"],
    ]
)

add_body(doc,
    "Le mécanisme principal de dégradation sous forte charge n'est pas le bruit thermique "
    "ou l'affaiblissement de trajet, mais la collision MAC en mode ALOHA pur : LoRaWAN "
    "n'implémente pas de mécanisme de collision avoidance (pas de CSMA). Deux nœuds "
    "émettant simultanément sur le même canal et le même SF produisent une collision "
    "irréversible. La fenêtre de collision est égale au Time-on-Air (51 ms pour SF7, "
    "1 318 ms pour SF12), ce qui explique pourquoi SF élevé aggrave significativement "
    "le taux de collision en réseau dense."
)

add_body(doc,
    "Ces résultats justifient directement les valeurs PDR choisies dans notre "
    "expérimentation. Pour un rucher isolé opérant avec 1 à 5 nœuds sur un gateway "
    "privé (configuration typique d'un apiculteur), le PDR nominal de 0,90 est "
    "conservateur : en pratique, on s'attendrait à PDR ≈ 0,95–0,99 en l'absence "
    "de voisins. La valeur PDR = 0,65 simule un environnement combinant forte "
    "atténuation par le feuillage (path-loss additionnel de 15–25 dB en été) et "
    "une présence éventuelle d'autres nœuds LoRaWAN partagés sur le même gateway "
    "communautaire (The Things Network, Helium)."
)

add_body(doc,
    "Un résultat complémentaire de LoRaSim concerne l'hétérogénéité des SF : "
    "une politique d'ADR bien configurée qui force les nœuds proches sur SF7 et "
    "les nœuds éloignés sur SF12 peut maintenir le PDR du réseau au-dessus de "
    "80 % jusqu'à N = 1 000 nœuds, contre un effondrement à 30 % avec SF12 "
    "homogène. Cette insight justifie l'importance de l'ADR (§ 2.3) pour toute "
    "architecture LoRaWAN à plus de 50 nœuds."
)

# ── 3.5 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.5  MQTT comme couche de transport locale", level=2)

add_body(doc,
    "Dans notre architecture, MQTT (Message Queuing Telemetry Transport, OASIS Standard "
    "v5.0) joue le rôle de couche de transport entre le simulateur biophysique "
    "(random_data_publisher.py) et la base de données time-series (mqtt_to_influx.py → "
    "InfluxDB). Ce choix est justifié par plusieurs propriétés du protocole adaptées "
    "à l'IoT temps réel :"
)

add_body(doc,
    "Architecture publish/subscribe découplée : le publisher et le subscriber ne "
    "connaissent pas leurs adresses respectives — seul le topic MQTT les relie. "
    "Cela permet de plugger un nouvel abonné (ex. un second algorithme de détection "
    "ou une alerte SMS) sans modifier le publisher, respectant le principe d'ouverture/"
    "fermeture de l'architecture logicielle."
)

add_body(doc,
    "Niveaux de QoS : MQTT définit trois niveaux de qualité de service. QoS 0 "
    "(At most once) est utilisé dans notre implémentation : le message est envoyé "
    "une fois sans acquittement. Cette configuration est intentionnelle — elle mime "
    "le comportement natif de LoRaWAN où les paquets perdus ne sont pas retransmis "
    "automatiquement (contrainte Duty Cycle). QoS 1 (At least once) garantirait la "
    "livraison au prix de doublons potentiels ; QoS 2 (Exactly once) évite les doublons "
    "mais quadruple le nombre d'échanges réseau — inadapté à un nœud LoRaWAN."
)

add_body(doc,
    "Légèreté du protocole : l'overhead MQTT est de 2 octets fixes + longueur du topic "
    "en entête fixe, soit ~20 octets pour notre topic `hive/ruche-01/telemetry`. "
    "Comparé aux 13+ octets d'overhead LoRaWAN (FHDR + MIC) sur un payload de 20 octets, "
    "la fraction utile reste comparable. En revanche, MQTT sur TCP introduit une latence "
    "de connexion (handshake TCP + CONNECT/CONNACK) qui, dans notre contexte local "
    "(broker Mosquitto sur le même réseau Docker), est négligeable (< 1 ms)."
)

add_table(doc,
    ["Critère", "MQTT (QoS 0)", "CoAP (UDP)", "AMQP 1.0"],
    [
        ["Transport",        "TCP",              "UDP",               "TCP"],
        ["Overhead entête",  "~20 B",            "4 B minimum",       "~8 B + framing"],
        ["Publish/Subscribe","Natif",            "Observe (RFC 7641)", "Natif"],
        ["Adapté IoT LPWAN", "Oui (QoS 0)",      "Optimal (UDP)",      "Non (lourd)"],
        ["Support TLS",      "Oui (MQTTS :8883)","Oui (DTLS)",         "Oui"],
        ["Usage projet",     "Mosquitto local",  "Non retenu",         "Non retenu"],
    ]
)

add_body(doc,
    "Pour une migration vers un déploiement LoRaWAN réel (The Things Network), "
    "le publisher MQTT serait remplacé par un nœud LoRa physique + gateway TTN. "
    "TTN expose nativement un API MQTT permettant de s'abonner aux uplinks des "
    "dispositifs via le topic `v3/{app_id}/devices/{dev_id}/up` — sans aucune "
    "modification requise sur mqtt_to_influx.py ni la suite du pipeline. Cette "
    "portabilité est l'un des avantages architecturaux majeurs du découplage MQTT."
)

# ── 3.6 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.6  Comparaison critique des technologies LPWAN", level=2)

add_body(doc,
    "Le choix de LoRaWAN comme protocole de référence pour ce projet s'inscrit dans "
    "un paysage LPWAN plus large. Le tableau suivant positionne LoRaWAN face à ses "
    "deux principaux concurrents académiques et industriels, sur les critères "
    "pertinents pour un déploiement apicole :"
)

add_table(doc,
    ["Critère", "LoRaWAN", "Sigfox", "NB-IoT (3GPP)"],
    [
        ["Bande fréquence",  "ISM libre (868 MHz EU)",   "ISM libre (868 MHz EU)",   "Licenciée (bandes LTE)"],
        ["Débit montant",    "0,3–5,5 kbps",             "100 bps (12 B/msg max)",   "20–127 kbps"],
        ["Portée rurale",    "2–15 km",                  "3–50 km",                  "1–10 km (indoor ++)"],
        ["Sécurité",         "AES-128 OTAA (bout-réseau)","AES-128 (réseau Sigfox)",  "LTE-grade (bout-bout)"],
        ["Msgs/jour max",    "~1 400 (DC 1 %)",          "140 uplink / 4 downlink",  "Illimité (abonnement)"],
        ["Coût infrastructure","Gateway ~100 €",         "Abonnement réseau",        "Opérateur télécom"],
        ["Déploiement privé","Oui (gateway TTN/perso)",  "Non (réseau fermé)",       "Non (réseau opérateur)"],
        ["Pertinence apicole","Optimal",                 "Limité (140 msg/j)",       "Trop coûteux, couverture rurale limitée"],
    ]
)

add_body(doc,
    "La capacité de déploiement d'un réseau LoRaWAN privé (gateway à ~100 € permettant "
    "de couvrir un rayon de 5–10 km) est un atout décisif pour les apiculteurs opérant "
    "dans des zones rurales sans couverture NB-IoT. La limite de 140 messages par jour "
    "de Sigfox (contre ~1 400 pour LoRaWAN en DC 1 %) exclut toute surveillance "
    "à intervalles inférieurs à 10 minutes, incompatible avec la détection d'essaimage "
    "qui requiert une résolution temporelle de 1 à 5 minutes."
)

# ---------------------------------------------------------------------------
# Sauvegarde
# ---------------------------------------------------------------------------
out_path = "Chapitre_3_LoRaWAN.docx"
doc.save(out_path)
print(f"✓ Document généré : {out_path}")
